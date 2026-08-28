from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(r"C:\Users\ASUS\OneDrive\Desktop\New folder\sitepilot\investor-materials")
ASSETS = ROOT / "assets"
OUT = ROOT / "OrchPilot-Investor-Strategy-and-Product-Brief.docx"

BLUE = "276EF1"; NAVY = "0B1736"; MUTED = "586174"; PALE = "EDF4FF"; PANEL = "F1F3F6"; WHITE = "FFFFFF"; GREEN = "2DA66F"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.82); sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.18
for name, size, color, before, after in [("Heading 1",18,BLUE,16,8),("Heading 2",14,NAVY,12,6),("Heading 3",11.5,MUTED,9,4)]:
    st=styles[name]; st.font.name="Aptos Display"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

if "Callout" not in styles:
    cs=styles.add_style("Callout",WD_STYLE_TYPE.PARAGRAPH)
else: cs=styles["Callout"]
cs.font.name="Aptos"; cs.font.size=Pt(11); cs.font.bold=True; cs.font.color.rgb=RGBColor.from_string(NAVY)
cs.paragraph_format.space_before=Pt(8); cs.paragraph_format.space_after=Pt(8); cs.paragraph_format.left_indent=Inches(0.18); cs.paragraph_format.right_indent=Inches(0.18)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def borders(table, color="D7DCE5", size="4"):
    tblPr=table._tbl.tblPr; tb=tblPr.find(qn("w:tblBorders"))
    if tb is None: tb=OxmlElement("w:tblBorders"); tblPr.append(tb)
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el=OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),size); el.set(qn("w:color"),color); tb.append(el)

def set_cell_text(cell, text, bold=False, color=NAVY, size=9.5):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); r=p.add_run(text); r.bold=bold; r.font.name="Aptos"; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,WHITE,9); shade(t.rows[0].cells[i],NAVY)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],str(v),False,NAVY,9); shade(cells[i], WHITE if ridx%2 else "F7F8FA")
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    borders(t); doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def add_bullet(text, level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2"); p.paragraph_format.space_after=Pt(4); p.add_run(text); return p

def add_callout(text, fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Inches(6.75)
    c=t.cell(0,0); shade(c,fill); set_cell_text(c,text,True,NAVY,10.5); c.margin_top=c.margin_bottom=Inches(0.10)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def add_picture(path, caption, width=6.75):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(str(path),width=Inches(width))
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after=Pt(8)
    r=cp.add_run(caption); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(MUTED)

def add_source(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(5)
    r=p.add_run("Source: "+text); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED); r.italic=True

# Header/footer
header=sec.header.paragraphs[0]; header.text="ORCHPILOT  /  INVESTOR STRATEGY & PRODUCT BRIEF"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: r.font.name="Aptos"; r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=footer.add_run("CONFIDENTIAL DRAFT  •  AUGUST 2026"); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(10)
r=p.add_run("ORCHPILOT"); r.font.name="Aptos Display"; r.font.size=Pt(18); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16)
r=p.add_run("The governed orchestration layer\nfor people, systems and AI"); r.font.name="Aptos Display"; r.font.size=Pt(34); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(35)
r=p.add_run("Investor strategy, product analysis, competitive positioning and 18-month roadmap"); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(MUTED)
doc.add_page_break()
doc.add_heading("Product at a glance",level=1)
add_picture(ASSETS/"03-workflow-designer.png","The working OrchPilot workflow designer: AI, deterministic flow and human steps in one canvas.",6.4)
add_callout("Status: working product foundation with deep engineering evidence; commercial traction, founder story, raise amount and customer ROI remain to be supplied before external circulation.","F4F6F9")
doc.add_page_break()

doc.add_heading("Executive summary",level=1)
doc.add_paragraph("OrchPilot is an extensible workflow execution platform designed to coordinate long-running work across software systems, people and AI agents. Its strongest investment thesis is not generic no-code automation. The credible wedge is governed orchestration for organizations where approvals, exceptions, audit evidence, deployment control and recovery are materially valuable.")
add_callout("Investor thesis: OrchPilot can occupy the governed middle between easy-to-use automation tools and code-first durable execution platforms.")
doc.add_heading("What is already credible",level=2)
for b in [
    "A real execution engine with immutable workflow versions, checkpointed state, retries, compensation, synchronous/asynchronous/scheduled/event-driven modes and cluster-aware claiming.",
    "Human tasks as first-class workflow participants, including forms, assignment, claiming, drafts, expiration, reassignment and auditable history.",
    "Runtime plugin loading: validated JARs stored in GridFS, loaded with isolated classloaders and registered as new node types without rebuilding the engine.",
    "Enterprise control primitives: RBAC, JWT authentication, Argon2id password hashing, encrypted secrets, security audit logs and server-side validation.",
    "AI agent nodes and provider management that keep adaptive reasoning inside deterministic workflow boundaries."
]: add_bullet(b)
doc.add_heading("What investors will still need",level=2)
for b in ["A narrow beachhead and five design partners with quantified process baselines.","Enterprise readiness proof: SSO/SCIM, high-availability benchmarks, deployment automation, observability, security review and support model.","A pricing hypothesis validated by willingness-to-pay—not just a feature comparison.","Founder-market-fit narrative, current traction, financing amount, runway plan and ownership/cap-table detail."]: add_bullet(b)

doc.add_heading("1. Product: what OrchPilot does",level=1)
doc.add_paragraph("A workflow in OrchPilot is a graph. Built-in nodes control the process; plugin nodes connect external capabilities. Publishing snapshots an immutable version, and each execution pins that version so later edits cannot silently change a run already in flight. This is especially valuable for approvals and regulated processes that may remain open for hours or days.")
add_picture(ASSETS/"01-workflows-dashboard.png","Workflow inventory and lifecycle management in the authenticated console.")
doc.add_heading("Core user journeys",level=2)
add_table(["User","Job to be done","Product surface"],[
    ["Workflow owner","Design, validate, publish and operate a process","Designer, versions, validation, executions"],
    ["Approver / operator","Claim, review and complete accountable human work","Task inbox, forms, drafts, history"],
    ["Platform administrator","Control users, plugins, secrets, AI providers and storage","Admin, plugin manager, secrets, settings"],
    ["Integration developer","Create reusable, governed capabilities","Plugin SDK, registry, schemas, permissions"]
],[1.25,3.0,2.5])

doc.add_heading("2. Architecture and engineering assessment",level=1)
doc.add_paragraph("The repository is a Maven multi-module system built with Java 17+, Spring Boot, MongoDB and Angular. The engine is deliberately integration-agnostic: plugin contracts live in a separate SDK, while the core resolves node types through a registry rather than integration-specific switch statements.")
add_table(["Evidence","Observed value","Investor meaning"],[
    ["Java production source","650 files / 80,084 lines","Substantial platform implementation; requires maintainability discipline"],
    ["Java tests","141 files / 21,783 lines; 73 core test classes","Meaningful test investment, but scale/HA benchmarks still needed"],
    ["REST surface","160 handler annotations / 27 controllers","Broad operational and administrative API surface"],
    ["Plugin modules","13","Demonstrates SDK breadth across SaaS, cloud, data and infrastructure"],
    ["Angular source","118 files / 28,595 lines","Full administration and workflow console, not API-only middleware"]
],[1.55,2.2,3.0])
add_source("Local source-only inventory, 26 August 2026. Generated targets and dependencies excluded.")
doc.add_heading("Architectural strengths",level=2)
for b in [
    "Plugin lifecycle is treated as an operational security boundary: validation, checksums, allowlisted hosts, scoped secrets, activation/deactivation and invocation history.",
    "Execution recovery persists state at node boundaries and uses ownership/heartbeat semantics to resume abandoned work.",
    "Idempotency keys include execution, node, plugin version and configuration fingerprint, reducing duplicate side effects after retry or crash.",
    "Workflow portability uses a signed/encrypted .orchpilot format and intentionally excludes secrets.",
    "Security design is explicit and documented, including password migration, token strategy, permission mapping and audit collections."
]: add_bullet(b)
doc.add_heading("Technical gaps to address before enterprise claims",level=2)
for b in [
    "Publish repeatable load, recovery, chaos and multi-node benchmarks with defined SLOs.",
    "Add SSO/SAML or OIDC federation, SCIM lifecycle management, external secret-manager integrations and tenant isolation if multi-tenant SaaS is planned.",
    "Harden the plugin sandbox beyond classloader isolation for untrusted third-party code; consider out-of-process execution for marketplace plugins.",
    "Formalize database migration/versioning, backup/restore drills, upgrade compatibility and disaster-recovery runbooks.",
    "Reduce source breadth risk through module ownership, architectural fitness tests and a published compatibility policy for the SDK."
]: add_bullet(b)

doc.add_heading("3. Product experience and governance",level=1)
add_picture(ASSETS/"04-user-governance.png","User governance with roles, account status and security-log access.")
doc.add_paragraph("The console already exposes the operational surfaces enterprise buyers expect: workflows, executions, tasks, forms, node types, plugins, secrets, events, AI providers, storage, users and groups. This breadth is an advantage in diligence, but the pitch should focus on the few workflows that create budget urgency.")
add_picture(ASSETS/"05-ai-providers.png","AI provider management: keys remain encrypted and are not returned to the browser or stored in workflows.")
add_callout("Product principle to preserve: AI should reason inside governed boundaries; deterministic workflow state, policy and human escalation should own the business outcome.")

doc.add_page_break()
doc.add_heading("4. Market timing",level=1)
doc.add_paragraph("Current market signals support an orchestration thesis. Gartner forecasts spending on consolidated business orchestration and automation technologies (BOAT) to grow 35% in 2025 to nearly $7 billion and cross $21 billion by 2029. A broader Gartner market map projects process-agnostic hyperautomation software at $119.2 billion by 2028, growing at 15.9% CAGR. These are category forecasts, not OrchPilot revenue projections.")
add_table(["Signal","Current evidence","Implication"],[
    ["Platform consolidation","BOAT spending projected above $21B by 2029","Buyers want a control layer, not another isolated bot"],
    ["AI enters core processes","Agentic systems increase nondeterminism and governance needs","Deterministic + adaptive orchestration becomes strategic"],
    ["Self-hosting remains relevant","Leading vendors sell self-managed enterprise tiers","Deployment control can be a commercial feature"],
    ["Human work persists","Major orchestration platforms treat tasks as first-class","Approvals and exceptions are a durable wedge"]
],[1.55,2.75,2.45])
add_source("Gartner BOAT forecast (May 2025); Gartner hyperautomation opportunity map (February 2025); Camunda, Temporal and n8n official product pages, accessed August 2026.")

doc.add_heading("5. Competitive landscape",level=1)
add_table(["Product","Primary strength","Where OrchPilot should differentiate"],[
    ["Zapier / Make","Accessibility and very broad SaaS connector ecosystems","Do not compete on connector count; win governed, long-running operations"],
    ["n8n","Flexible visual automation, self-hosting and technical-user adoption","Stronger plugin isolation, human accountability and execution governance"],
    ["Camunda","Enterprise process orchestration, BPMN, governance and scale","Faster onboarding and a simpler plugin-native developer/operator experience"],
    ["Temporal","Category-leading durable execution for developers","Business-visible workflow design, human tasks and schema-driven integrations"],
    ["Workato / enterprise iPaaS","Enterprise integrations, governance and services ecosystem","Deployment control, developer extensibility and transparent execution economics"]
],[1.35,2.5,2.9])
doc.add_paragraph("The defensible positioning is a governed orchestration control plane that is easier to adopt than BPMN-heavy suites, more operationally accountable than general automation tools, and more business-visible than code-only durable execution.")

doc.add_heading("6. Beachhead, impact and go-to-market",level=1)
doc.add_heading("Recommended beachhead",level=2)
doc.add_paragraph("Target platform engineering, regulated operations and IT service-delivery teams in 500–5,000 employee organizations. These teams have hybrid systems, recurring approval-heavy processes, explicit compliance owners and enough failure cost to pay for orchestration reliability.")
add_table(["Beachhead workflow","Why it fits","Measure before / after"],[
    ["Cloud access & provisioning","Plugins + approvals + rollback + audit","Lead time, failed changes, manual touches"],
    ["Customer onboarding / KYC exception","Forms + human judgment + immutable history","Cycle time, rework, SLA breaches"],
    ["Incident remediation","Events + automation + escalation + evidence","MTTR, repeat incidents, operator time"],
    ["Vendor / security review","Long-running tasks + documents + approvals","Days to decision, overdue tasks, audit prep"]
],[1.7,2.45,2.6])
doc.add_heading("Why this can create market impact",level=2)
for b in [
    "Turn fragile tribal processes into explicit, versioned operational assets.",
    "Let organizations adopt AI for exceptions without surrendering auditability or deterministic control.",
    "Give internal platform teams a reusable self-service layer instead of bespoke portals and scripts for every request.",
    "Create a plugin supply chain where partners package domain operations once and customers govern their use centrally."
]: add_bullet(b)
doc.add_heading("Go-to-market motion",level=2)
for b in [
    "Founder-led discovery: interview 25 process owners and platform leaders; score pain, frequency, compliance and budget ownership.",
    "Five design partners: implement one workflow each under a fixed success plan with baseline metrics.",
    "Convert three paid pilots: annual platform commitment plus scoped launch package; price before building custom breadth.",
    "Publish proof: architecture/security brief, reliability benchmark and two quantified case studies.",
    "Build channel leverage: certify specialist integrators and plugin authors around one vertical pack."
]: add_bullet(b)

doc.add_heading("7. Business model recommendation",level=1)
add_table(["Revenue stream","Recommended structure","Validation question"],[
    ["Platform subscription","Annual contract by environment / deployment tier","Will governance and support unlock budget?"],
    ["Execution bands","Committed annual executions with transparent overage","Does usage track customer value predictably?"],
    ["Plugin marketplace","Certified paid plugins and private enterprise catalogues","Will partners build and customers pay for governed supply?"],
    ["Launch services","Fixed-scope implementation and enablement","Can services produce repeatable templates rather than custom debt?"]
],[1.45,3.0,2.3])
add_callout("Treat pricing as a testable hypothesis. Do not publish tiers until five design-partner conversations reveal procurement anchors and usage patterns.")

doc.add_page_break()
doc.add_heading("8. Eighteen-month product roadmap",level=1)
add_table(["Phase","Outcome","Product priorities","Evidence gate"],[
    ["0–3 months","Prove the wedge","Install path, first-use UX, vertical workflow pack, discovery instrumentation","5 design partners; one process baseline each"],
    ["4–9 months","Earn enterprise trust","SSO/SCIM, HA topology, observability, external secrets, backup/restore, policy controls","3 paid pilots; documented SLO and recovery test"],
    ["10–18 months","Scale distribution","Plugin certification, partner portal, SDK compatibility, cloud option, usage metering","Expansion to second process; partner-sourced pipeline"]
],[0.85,1.35,3.2,1.35])
doc.add_heading("North-star metrics",level=2)
for b in ["Time from install to first production workflow.","Successful workflow completion rate and recovery rate after induced failure.","Human wait time, overdue-task rate and exception resolution time.","Number of production workflows per customer and expansion to a second department/process.","Certified plugins used in production and partner-sourced annual recurring revenue."]: add_bullet(b)

doc.add_heading("9. Risks and mitigation",level=1)
add_table(["Risk","Why it matters","Mitigation"],[
    ["Crowded positioning","Automation language collapses into larger incumbent categories","Own governed orchestration for high-accountability workflows"],
    ["Enterprise readiness gap","Feature breadth does not equal procurement readiness","Security roadmap, benchmarks, SLOs, support and deployment proof"],
    ["Plugin security","In-process third-party code can exceed classloader boundaries","Certification, signing, policy and optional out-of-process runtime"],
    ["Connector cold start","Customers expect their systems on day one","Vertical packs, REST fallback and partner-led plugin supply"],
    ["Breadth before PMF","Large codebase can absorb effort without customer learning","Freeze non-wedge features until design-partner evidence"]
],[1.25,2.55,3.1])

doc.add_heading("10. Investor pitch guidance",level=1)
doc.add_heading("Thirty-second version",level=2)
add_callout("OrchPilot is the governed orchestration layer for critical workflows that span people, systems and AI. Unlike task automation tools, it owns long-running state, human accountability and recovery. Unlike code-only engines, it gives operators a visual, extensible control plane. The product foundation is built; the seed round converts it into repeatable enterprise adoption.")
doc.add_heading("Questions to answer before pitching externally",level=2)
for b in [
    "Who are the founders, and what unique access or insight makes this team credible in platform engineering or regulated operations?",
    "What is the raise amount, target runway, hiring plan and milestone-based allocation?",
    "Which design partners or conversations exist today, and what process pain has been quantified?",
    "What is the deployment model and license? Open core, commercial source, hosted SaaS or dual model?",
    "What is the first vertical plugin pack, and why can OrchPilot win distribution there?",
    "What scale has been tested: concurrent executions, workflow duration, recovery time and plugin isolation under failure?"
]: add_bullet(b)

doc.add_heading("Appendix A. Product capability map",level=1)
add_table(["Domain","Implemented evidence","Pitch status"],[
    ["Workflow lifecycle","Draft, validate, publish, immutable versions, archive, import/export","Built"],
    ["Execution","Sync/async/manual/event/schedule, retry, compensation, pause/resume/cancel","Built"],
    ["Human work","Forms, tasks, assignment, claim, draft, completion, history, expiry","Built"],
    ["Plugins","SDK, registry client/server, upload, validation, activation, isolation, history","Built; marketplace is roadmap"],
    ["Security","RBAC, JWT, Argon2id, audit, encrypted secrets, CSP/security headers","Built; enterprise identity is roadmap"],
    ["AI","Agent node, provider configuration, usage and CLI controls","Built foundation; production proof needed"],
    ["Operations","Health, metrics, logs, schedulers, recovery, cluster claims","Built foundation; benchmarks needed"]
],[1.25,4.0,1.65])

doc.add_heading("Appendix B. Sources",level=1)
sources=[
    "Workflow-OrchPilot/workflow-engine local repository: README.md, ARCHITECTURE.md, SECURITY.md, AI_AGENT_NODE.md, AI_CLI_CONFIGURATION.md, WORKFLOW_INSTANCE_LIFECYCLE.md, WORKFLOW_PORTABILITY.md and PLUGIN_OPERATION_FRAMEWORK.md (reviewed 26 August 2026).",
    "Gartner, Forecast Analysis: Business Orchestration and Automation Technologies, Worldwide (28 May 2025): https://www.gartner.com/en/documents/6530402",
    "Gartner, Market Opportunity Map: Hyperautomation via Process-Agnostic Software, Worldwide (11 February 2025): https://www.gartner.com/en/documents/6165623",
    "Gartner newsroom, enterprise network automation forecast (18 September 2024): https://www.gartner.com/en/newsroom/press-releases/2024-09-18-gartner-says-30-percent-of-enterprises-will-automate-more-than-half-of-their-network-activities-by-2026",
    "Camunda platform and use cases: https://camunda.com/platform/ and https://camunda.com/platform/use-cases/",
    "Temporal platform and platform-engineering solution: https://temporal.io/ and https://temporal.io/solutions/platform-engineering",
    "n8n pricing and Business plan: https://n8n.io/pricing/ and https://support.n8n.io/article/what-is-the-n-8-n-business-plan-that-launched-in-august-2025"
]
for s in sources:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.18); p.paragraph_format.first_line_indent=Inches(-0.12); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.0
    r=p.add_run("•  "+s); r.font.name="Aptos"; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(NAVY)

doc.core_properties.title="OrchPilot Investor Strategy and Product Brief"
doc.core_properties.subject="Investor pitch, codebase assessment, market analysis and product roadmap"
doc.core_properties.author="OrchPilot"
doc.save(OUT)
print(OUT)
